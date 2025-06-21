import os
import unittest
import tempfile
import shutil
from unittest.mock import patch

# Importar las funciones a testear
from certificados_utils import (
    obtener_curso_completo,
    fpuntoscedula,
    formatear_fecha,
    obtener_parametro
)
from generar_certificados import editardatos

class TestGenerarCertificados(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        # Crear un directorio temporal para las pruebas
        cls.test_dir = tempfile.mkdtemp()
        cls.docx_dir = os.path.join(cls.test_dir, "docx")
        cls.pdf_dir = os.path.join(cls.test_dir, "pdf")
        os.makedirs(cls.docx_dir, exist_ok=True)
        os.makedirs(cls.pdf_dir, exist_ok=True)
        
        # Crear una plantilla simple de prueba
        cls.template_path = os.path.join(cls.test_dir, "plantilla_test.docx")
        cls.create_test_template(cls.template_path)
        
        # Configurar datos de prueba
        cls.test_data = {
            "nombre": "JUAN PEREZ",
            "numerocedula": "12345678",
            "curse": "Humanización",
            "date": "2023-05-15 10:00:00",
            "cargo": "ENFERMERO",
            "ri": "TEST123",
            "output_file": "certificado_test.docx"
        }

    @classmethod
    def tearDownClass(cls):
        # Eliminar el directorio temporal después de las pruebas
        shutil.rmtree(cls.test_dir)

    @staticmethod
    def create_test_template(template_path):
        """Crea una plantilla de Word simple para pruebas"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        
        # Añadir texto con los campos que espera la función editardatos
        p = doc.add_paragraph()
        p.add_run("Nombre: ").bold = True
        p.add_run("<<nombre>>")
        
        p = doc.add_paragraph()
        p.add_run("Documento: ").bold = True
        p.add_run("C.C. <<cedula>>")
        
        p = doc.add_paragraph()
        p.add_run("Fecha: ").bold = True
        p.add_run("<<fecha>>")
        
        p = doc.add_paragraph()
        p.add_run("Curso: ").bold = True
        p.add_run("<<curso>>")
        
        p = doc.add_paragraph()
        p.add_run("Cargo: ").bold = True
        p.add_run("<<cargo>>")
        
        p = doc.add_paragraph()
        p.add_run("RI: ").bold = True
        p.add_run("<<ri>>")
        
        doc.save(template_path)

    def test_obtener_curso_completo(self):
        """Test para verificar la función obtener_curso_completo"""
        self.assertEqual(
            obtener_curso_completo("Humanización"),
            "HUMANIZACIÓN DE LA ATENCIÓN EN SALUD"
        )
        self.assertEqual(
            obtener_curso_completo("Curso no existente"),
            "Curso no existente"
        )

    def test_fpuntoscedula(self):
        """Test para verificar la función fpuntoscedula"""
        self.assertEqual(fpuntoscedula("12345678"), "12.345.678")
        self.assertEqual(fpuntoscedula("123"), "123")
        self.assertEqual(fpuntoscedula("123456789"), "123.456.789")

    def test_formatear_fecha(self):
        """Test para verificar la función formatear_fecha"""
        self.assertEqual(
            formatear_fecha("2023-05-15 10:00:00"),
            "15 DE MAYO DE 2023"
        )
        self.assertEqual(
            formatear_fecha("15/05/2023"),
            "15 DE MAYO DE 2023"
        )
        # Test para fecha inválida
        self.assertEqual(
            formatear_fecha("fecha inválida"),
            "fecha inválida"
        )

    def test_obtener_parametro(self):
        """Test para verificar la función obtener_parametro"""
        self.assertTrue(
            len(obtener_parametro("HUMANIZACIÓN DE LA ATENCIÓN EN SALUD")) > 10
        )
        self.assertEqual(
            obtener_parametro("Curso no existente"),
            "-"
        )

    @patch('generar_certificados.TEMPLATE_PATH', new_callable=lambda: TestGenerarCertificados.template_path)
    @patch('generar_certificados.DOCX_DIR', new_callable=lambda: TestGenerarCertificados.docx_dir)
    @patch('generar_certificados.PDF_DIR', new_callable=lambda: TestGenerarCertificados.pdf_dir)
    def test_editardatos(self, *args):
        """Test para verificar la función editardatos con un solo dato"""
        # Ejecutar la función con los datos de prueba
        editardatos(
            self.test_data["nombre"],
            self.test_data["numerocedula"],
            self.test_data["curse"],
            self.test_data["date"],
            self.test_data["cargo"],
            self.test_data["ri"],
            self.test_data["output_file"]
        )
        
        # Verificar que se creó el archivo DOCX
        docx_path = os.path.join(self.docx_dir, self.test_data["output_file"])
        self.assertTrue(os.path.exists(docx_path), "El archivo DOCX no se creó correctamente")
        
        # Verificar que se creó el archivo PDF (opcional, dependiendo de si tienes docx2pdf instalado)
        pdf_path = os.path.join(self.pdf_dir, self.test_data["output_file"].replace(".docx", ".pdf"))
        # self.assertTrue(os.path.exists(pdf_path), "El archivo PDF no se creó correctamente")

if __name__ == "__main__":
    unittest.main()