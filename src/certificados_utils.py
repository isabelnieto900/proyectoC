import re
from datetime import datetime
import os
import pandas as pd
from docx import Document
from pathlib import Path

mapeo_cursos = {
    "Humanización": "HUMANIZACIÓN DE LA ATENCIÓN EN SALUD",
    "Gestión del duelo": "GESTIÓN Y MANEJO DEL DUELO EN LOS SERVICIOS DE SALUD",
    "Violencia sexual": "ATENCIÓN INTEGRAL EN SALUD A VÍCTIMAS DE VIOLENCIA SEXUAL",
    "Agentes químicos": "ATENCIÓN INTEGRAL EN SALUD A VÍCTIMAS DE ATAQUES CON AGENTES QUÍMICOS",
    "Soporte vital básico": "SOPORTE VITAL BÁSICO",
    "Soporte vital avanzado": "SOPORTE VITAL AVANZADO",
    "Cuidado del donante": "DETECCIÓN Y CUIDADO DEL DONANTE DE ÓRGANOS Y TEJIDOS",
    # Agrega más mapeos según sea necesario
}

def obtener_curso_completo(curse):
    return mapeo_cursos.get(curse, curse)

def fpuntoscedula(numerocedula):
    numerocedula = str(numerocedula)
    # Agrupa de a tres dígitos desde la derecha
    partes = []
    while numerocedula:
        partes.insert(0, numerocedula[-3:])
        numerocedula = numerocedula[:-3]
    cedula2 = ".".join(partes)
    return cedula2

fpuntoscedula("79619940")
def formatear_fecha(fecha_str):
    try:
        fecha_obj = datetime.strptime(fecha_str, "%Y-%m-%d %H:%M:%S")
    except ValueError:
        try:
            fecha_obj = datetime.strptime(fecha_str, "%d/%m/%Y")
        except ValueError as e:
            print(f"❌ Error al formatear la fecha {fecha_str}: {e}")
            return fecha_str
            exit
    return fecha_obj.strftime("%d DE %B DE %Y").upper()

def obtener_parametro(curse_completo):
    if curse_completo in ["SOPORTE VITAL BÁSICO", "SOPORTE VITAL AVANZADO"]:
        return "BAJO CONSENSOS ILCOR 2020, RECOMENDACIONES Y ACTUALIZACIONES AHA 2020, Y LO REGLAMENTADO EN LA RESOLUCIÓN 3100 DE 2019 DEL MINISTERIO DE SALUD Y PROTECCIÓN SOCIAL."
    elif curse_completo == "DETECCIÓN Y CUIDADO DEL DONANTE DE ÓRGANOS Y TEJIDOS":
        return "SEGÚN LO REGLAMENTADO EN LA LEY 1805 DE 2016, RESOLUCIÓN 3100 DE 2019 DEL MINISTERIO DE SALUD Y PROTECCIÓN SOCIAL, RESOLUCIÓN 0156 DE 2021 Y RESOLUCIÓN 0317 DE 2022 DEL INSTITUTO NACIONAL DE SALUD."
    elif curse_completo == "GESTIÓN Y MANEJO DEL DUELO EN LOS SERVICIOS DE SALUD":
        return "SEGÚN LO REGLAMENTADO EN LA RESOLUCIÓN 3100 DE 2019 DEL MINISTERIO DE SALUD Y DE LA PROTECCIÓN SOCIAL."
    elif curse_completo == "HUMANIZACIÓN DE LA ATENCIÓN EN SALUD":
        return "BASADO EN EL PLAN NACIONAL DE MEJORAMIENTO DE LA CALIDAD EN SALUD Y EN LA POLÍTICA NACIONAL DE HUMANIZACIÓN EN SALUD DEL MINISTERIO DE SALUD Y PROTECCIÓN SOCIAL, Y LO REGLAMENTADO EN LA LEY 1438 DE 2011"
    elif curse_completo == "ATENCIÓN INTEGRAL EN SALUD A VÍCTIMAS DE VIOLENCIA SEXUAL":
        return "BAJO PARÁMETROS ESTABLECIDOS EN LA LEY 1146 DE 2007 Y LA RESOLUCIÓN 0459 DE 2012, Y LO REGLAMENTADO EN LA RESOLUCIÓN 3100 DE 2019 DEL MINISTERIO DE SALUD Y PROTECCIÓN SOCIAL."
    elif curse_completo == "ATENCIÓN INTEGRAL EN SALUD A VÍCTIMAS DE ATAQUES CON AGENTES QUÍMICOS":
        return "EN CUMPLIMIENTO DE LA LEY 1971 DE 2019, PARÁMETROS ESTABLECIDOS EN LA RESOLUCIÓN 4568 DE 2014, Y LO REGLAMENTADO EN LA RESOLUCIÓN 3100 DE 2019 DEL MINISTERIO DE SALUD Y PROTECCIÓN SOCIAL"
    return "-"

class CertificadoGenerator:
    def __init__(self, template_path: str, output_dir: str):
        """
        Inicializa el generador de certificados
        
        Args:
            template_path: Ruta a la plantilla .docx
            output_dir: Directorio donde guardar los certificados generados
        """
        self.template_path = Path(template_path)
        self.output_dir = Path(output_dir)
        
        if not self.template_path.exists():
            raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")
        
        # Crear directorio de salida si no existe
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def replace_placeholders(self, doc: Document, replacements: dict):
        """
        Reemplaza placeholders en el documento
        
        Args:
            doc: Documento de Word
            replacements: Diccionario con los reemplazos {placeholder: valor}
        """
        # Reemplazar en párrafos
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
        
        # Reemplazar en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
    
    def generate_certificate(self, data: dict, filename: str):
        """
        Genera un certificado individual
        
        Args:
            data: Datos para reemplazar en la plantilla
            filename: Nombre del archivo de salida
        """
        # Cargar plantilla
        doc = Document(self.template_path)
        
        # Reemplazar placeholders
        self.replace_placeholders(doc, data)
        
        # Guardar certificado
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        return output_path
    
    def generate_from_excel(self, excel_path: str):
        """
        Genera certificados desde un archivo Excel
        
        Args:
            excel_path: Ruta al archivo Excel con los datos
        """
        try:
            df = pd.read_excel(excel_path)
            print(f"📊 Procesando {len(df)} registros...")
            
            for index, row in df.iterrows():
                # Convertir la fila a diccionario para los reemplazos
                data = row.to_dict()
                
                # Generar nombre de archivo (usar primera columna como identificador)
                first_column = df.columns[0]
                identifier = str(data[first_column]).replace(" ", "_")
                filename = f"certificado_{identifier}.docx"
                
                # Generar certificado
                output_path = self.generate_certificate(data, filename)
                print(f"✅ Generado: {filename}")
            
            print(f"🎉 Proceso completado. {len(df)} certificados generados en {self.output_dir}")
            
        except Exception as e:
            print(f"❌ Error procesando Excel: {e}")
            raise